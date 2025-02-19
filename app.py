import streamlit as st
import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import re

# Función para limpiar Markdown
def clean_markdown(text):
    """Elimina marcas de Markdown del texto."""
    text = re.sub(r'[#*_`]', '', text)  # Eliminar caracteres especiales de Markdown
    return text.strip()

# Función para procesar diálogos y listas, reemplazando guiones por rayas
def process_dialogues_and_lists(text):
    """
    Procesa el texto para:
    1. Reemplazar guiones ('-') al inicio de las listas o diálogos por rayas ('—').
    2. Asegurar que después de las listas haya un salto de párrafo.
    """
    lines = text.split('\n')  # Divide el texto en líneas
    processed_lines = []
    in_list = False  # Indicador para saber si estamos dentro de una lista o diálogo

    for line in lines:
        stripped_line = line.strip()
        if stripped_line.startswith('-'):  # Detectar líneas que comienzan con un guion
            processed_line = stripped_line.replace('-', '—', 1)
            processed_lines.append(processed_line)
            in_list = True
        else:
            if in_list:
                processed_lines.append("")  # Salto de párrafo
                in_list = False
            processed_lines.append(stripped_line)

    return '\n\n'.join(processed_lines)

# Función para aplicar reglas de capitalización según el idioma
def format_title(title, language):
    """
    Formatea el título según las reglas gramaticales del idioma.
    - Español: Solo mayúscula inicial en la primera palabra y nombres propios.
    - Otros idiomas: Mayúscula inicial en cada palabra.
    """
    if language.lower() == "spanish":
        words = title.split()
        formatted_words = [words[0].capitalize()] + [word.lower() if word.islower() else word for word in words[1:]]
        return " ".join(formatted_words)
    else:
        return title.title()

# Función para generar un capítulo usando OpenRouter AI
def generate_chapter(api_key, topic, audience, chapter_number, language, table_of_contents="", specific_instructions=""):
    url = "https://openrouter.ai/api/v1/chat/completions"
    
    # Construir el mensaje con la tabla de contenido e instrucciones específicas
    message_content = f"Escribe el capítulo {chapter_number} sobre {topic} dirigido a {audience}."
    
    if table_of_contents:
        message_content += f" Sigue esta estructura: {table_of_contents}"
    
    if specific_instructions:
        message_content += f" {specific_instructions}"
    
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    
    data = {
        "model": "sophosympatheia/rogue-rose-103b-v0.2:free",
        "messages": [
            {
                "role": "user",
                "content": message_content
            }
        ]
    }
    
    try:
        response = requests.post(url, headers=headers, json=data)
        response.raise_for_status()
        content = response.json().get("choices", [{}])[0].get("message", {}).get("content", "Error generating the chapter.")
    except requests.RequestException as e:
        st.error(f"Error al generar el capítulo {chapter_number}: {str(e)}")
        return "Error al generar el capítulo."
    
    # Procesar diálogos y listas
    processed_content = process_dialogues_and_lists(content)
    
    return clean_markdown(processed_content)

# Función para agregar numeración de páginas al documento Word
def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldChar2)

# Función para crear un documento Word con formato específico
def create_word_document(chapters, title, author_name, author_bio, language):
    doc = Document()

    # Configurar el tamaño de página (5.5 x 8.5 pulgadas)
    section = doc.sections[0]
    section.page_width = Inches(5.5)
    section.page_height = Inches(8.5)

    # Configurar márgenes de 0.8 pulgadas en todo
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # Añadir y formatear el título
    formatted_title = format_title(title, language)
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(formatted_title)
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = "Times New Roman"

    # Añadir nombre del autor si está proporcionado
    if author_name:
        author_paragraph = doc.add_paragraph()
        author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_paragraph.add_run(author_name)
        author_run.font.size = Pt(12)
        author_run.font.name = "Times New Roman"
        doc.add_page_break()

    # Añadir perfil del autor si está proporcionado
    if author_bio:
        bio_paragraph = doc.add_paragraph("Author Information")
        bio_paragraph.style = "Heading 2"
        bio_paragraph.runs[0].font.size = Pt(11)
        bio_paragraph.runs[0].font.name = "Times New Roman"
        doc.add_paragraph(author_bio).style = "Normal"
        doc.add_page_break()

    # Añadir capítulos
    for i, chapter in enumerate(chapters, 1):
        chapter_title_text = f"Chapter {i}" if language.lower() != "spanish" else f"Capítulo {i}"
        formatted_chapter_title = format_title(chapter_title_text, language)
        chapter_title = doc.add_paragraph(formatted_chapter_title)
        chapter_title.style = "Heading 1"
        chapter_title.runs[0].font.size = Pt(12)
        chapter_title.runs[0].font.name = "Times New Roman"

        paragraphs = chapter.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.replace('\n', ' ').strip()
            paragraph = doc.add_paragraph(para_text)
            paragraph.style = "Normal"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.name = "Times New Roman"

        doc.add_page_break()

    add_page_numbers(doc)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Configuración de Streamlit
st.set_page_config(page_title="Automatic Book Generator", page_icon="📚")

# Título con ícono
st.title("📚 Automatic Book Generator")

# Barra lateral con instrucciones y anuncio
st.sidebar.header("📖 ¿Cómo funciona esta aplicación?")
st.sidebar.markdown("""
Esta aplicación genera automáticamente libros en formato `.docx` basados en un tema y audiencia objetivo. 
Los libros pueden ser **ficción** o **no-ficción**, dependiendo de tu entrada.

**Pasos para usarla:**
1. Introduce el tema del libro.
2. Especifica la audiencia objetivo.
3. Proporciona una tabla de contenidos opcional.
4. Escribe instrucciones específicas opcionales.
5. Selecciona el número de capítulos deseado (máximo 50).
6. Elige el idioma del libro.
7. Decide si incluir nombre del autor y perfil del autor.
8. Haz clic en "Generar Libro".
9. Descarga el archivo generado.
""")
st.sidebar.markdown("""
---
**📝 Corrección de texto en 24 horas**  
👉 [Hablemos Bien](https://hablemosbien.org)
""")

# Validación de claves secretas
if "OPENROUTER_API_KEY" not in st.secrets:
    st.error("Por favor, configura la clave API en los secretos de Streamlit.")
    st.stop()
api_key = st.secrets["OPENROUTER_API_KEY"]

# Entradas del usuario
topic = st.text_input("📒 Tema del libro:")
audience = st.text_input("🎯 Audiencia objetivo:")
table_of_contents = st.text_area("📚 Tabla de contenidos opcional:", placeholder="Proporciona una tabla de contenidos para capítulos más largos.")
specific_instructions = st.text_area("📝 Instrucciones específicas opcionales:", placeholder="Proporciona instrucciones específicas para el libro.")
num_chapters = st.slider("🔢 Número de Capítulos", min_value=1, max_value=50, value=25)
author_name = st.text_input("🖋️ Nombre del Autor (opcional):")
author_bio = st.text_area("👤 Perfil del Autor (opcional):", placeholder="Descripción profesional breve o biografía.")
languages = ["English", "Spanish", "French", "German", "Chinese", "Japanese", "Russian", "Portuguese", "Italian", "Arabic", "Medieval Latin", "Koine Greek"]
selected_language = st.selectbox("🌐 Elige el idioma del libro:", languages)

# Estado de Streamlit para almacenar los capítulos generados
if 'chapters' not in st.session_state:
    st.session_state.chapters = []

# Botón para generar el libro
if st.button("🚀 Generar Libro"):
    if not topic or not audience:
        st.error("Por favor, introduce un tema y una audiencia objetivo válidos.")
        st.stop()

    chapters = []

    # Generar capítulos principales
    progress_bar = st.progress(0)
    for i in range(1, num_chapters + 1):
        st.write(f"⏳ Generando capítulo {i}...")
        chapter_content = generate_chapter(api_key, topic, audience, i, selected_language.lower(), table_of_contents, specific_instructions)
        word_count = len(chapter_content.split())
        chapters.append(chapter_content)
        with st.expander(f"📖 Capítulo {i} ({word_count} palabras)"):
            st.write(chapter_content)
        progress_bar.progress(i / num_chapters)

    st.session_state.chapters = chapters

# Mostrar opciones de descarga si hay capítulos generados
if st.session_state.chapters:
    st.subheader("⬇️ Opciones de Descarga")
    word_file = create_word_document(st.session_state.chapters, topic, author_name, author_bio, selected_language.lower())

    st.download_button(
        label="📥 Descargar en Word",
        data=word_file.getvalue(),
        file_name=f"{topic}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
